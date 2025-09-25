#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
cdf13_report.py — build an A4 Word report from cdf13_post outputs.

Usage
  python cdf13_report.py kh18.cdf
  python cdf13_report.py kh18.cdf --img kh18_with_baseline.png --peaks kh18_peaks.csv --out kh18_report.docx

Behavior
- Infers defaults from the CDF name:
    <base>_with_baseline.png   (figure from cdf13_post.py)
    <base>_peaks.csv           (integration table)
- Creates a single-page A4 .docx:
    • Top half: the figure, scaled to fit width and ≤ half the usable height
    • Bottom half: the peaks table, all columns found in the CSV
Deps
  pip install python-docx pandas pillow
"""

import argparse, os, sys
import pandas as pd
from PIL import Image
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def infer_paths(cdf_path, img_arg, peaks_arg, out_arg):
    base = os.path.splitext(os.path.basename(cdf_path))[0]
    img = img_arg or f"{base}_with_baseline.png"
    peaks = peaks_arg or f"{base}_peaks.csv"
    out = out_arg or f"{base}_report.docx"
    return img, peaks, out

def ensure_exists(path, kind):
    if not os.path.exists(path):
        sys.exit(f"ERROR: {kind} not found: {path}")

def set_page_a4(section, margin_mm=20):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    m = Mm(margin_mm)
    section.left_margin = m
    section.right_margin = m
    section.top_margin = m
    section.bottom_margin = m

def add_top_image(doc, img_path, section, half_ratio=0.5):
    # Usable area
    usable_w = section.page_width - section.left_margin - section.right_margin
    usable_h = section.page_height - section.top_margin - section.bottom_margin
    target_h = usable_h * half_ratio

    # Image aspect and scaling: prefer height cap, fallback to width cap if needed
    with Image.open(img_path) as im:
        w_px, h_px = im.size
    # Word uses EMU via Mm; we only need ratios
    scaled_w_by_h = (w_px / h_px) * target_h
    if scaled_w_by_h <= usable_w:
        # Set height only to keep ≤ half page
        pic = doc.add_picture(img_path, height=target_h)
    else:
        # Too wide at that height → cap by width
        pic = doc.add_picture(img_path, width=usable_w)

    # Center the image
    last_par = doc.paragraphs[-1]
    last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

def df_from_csv(csv_path):
    df = pd.read_csv(csv_path)
    # Keep column order as in file
    return df

def add_table(doc, df):
    # Basic Normal style tweaks
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    rows, cols = df.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.autofit = True

    # Header
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        run = cell.paragraphs[0].add_run(str(col))
        run.bold = True

    # Body
    for i in range(rows):
        for j in range(cols):
            val = df.iat[i, j]
            table.cell(i + 1, j).text = "" if pd.isna(val) else str(val)

def main():
    ap = argparse.ArgumentParser(description="Create A4 DOCX with plot on top and integration table below.")
    ap.add_argument("cdf", help="Original .cdf filename (used only for base name inference)")
    ap.add_argument("--img", help="Path to figure PNG/SVG; default <base>_with_baseline.png")
    ap.add_argument("--peaks", help="Path to peaks CSV; default <base>_peaks.csv")
    ap.add_argument("--out", help="Output DOCX; default <base>_report.docx")
    ap.add_argument("--margin-mm", type=float, default=20.0, help="Page margins in mm (all sides); default 20")
    args = ap.parse_args()

    img_path, peaks_path, out_path = infer_paths(args.cdf, args.img, args.peaks, args.out)
    ensure_exists(img_path, "figure")
    ensure_exists(peaks_path, "peaks CSV")

    df = df_from_csv(peaks_path)

    doc = Document()
    set_page_a4(doc.sections[0], margin_mm=args.margin_mm)

    # Top half: image
    add_top_image(doc, img_path, doc.sections[0], half_ratio=0.5)

    # Spacer
    doc.add_paragraph("")  # minimal spacing

    # Bottom half: table
    add_table(doc, df)

    doc.save(out_path)
    print(f"Wrote: {out_path}")

if __name__ == "__main__":
    main()
