# scripts/processing.py
import time
import os
import math
import sys
import shutil
import shlex
from pathlib import Path
import subprocess
import uuid
import numpy as np
import pandas as pd
from jinja2 import Environment, FileSystemLoader
from PIL import Image
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define project paths
PROJECT_ROOT = Path(__file__).parent.parent
REPORTS_DIR = PROJECT_ROOT / "generated_reports"
ASSETS_DIR = PROJECT_ROOT / "assets"
SCRIPTS_DIR_PATH = PROJECT_ROOT / "scripts"
REPORTS_DIR.mkdir(exist_ok=True)
ASSETS_DIR.mkdir(exist_ok=True)

# Replicate presets from cdf13_run.py to expand them into arguments for cdf13_post3.py
PRESETS = {
    "default": {"noise-mode":"pre","noise-window":"0.3","snr":"2","smooth":"8","min-width-sec":"1","measure":"area","reject-height":"0"},
    "lenient": {"noise-mode":"pre","noise-window":"0.5","snr":"1.5","smooth":"5","min-width-sec":"0.5","measure":"area","reject-height":"0"},
    "strict":  {"noise-mode":"pre","noise-window":"0.3","snr":"3","smooth":"16","min-width-sec":"2","measure":"area","reject-height":"0"},
}

# --- DOCX Generation Helpers (adapted from cdf13_report2.py) ---

def set_page_a4(section, margin_mm=20):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    m = Mm(margin_mm)
    section.left_margin, section.right_margin, section.top_margin, section.bottom_margin = m, m, m, m

def add_top_image(doc, img_path, section, half_ratio=0.5):
    usable_w = section.page_width - section.left_margin - section.right_margin
    usable_h = section.page_height - section.top_margin - section.bottom_margin
    target_h = usable_h * half_ratio
    with Image.open(img_path) as im:
        w_px, h_px = im.size
    scaled_w_by_h = (w_px / h_px) * target_h
    pic = doc.add_picture(img_path, height=target_h if scaled_w_by_h <= usable_w else None, width=usable_w if scaled_w_by_h > usable_w else None)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def decompose_sci(val):
    try:
        v = float(val)
        if v == 0.0: return ("0.00", 0)
        exp = int(math.floor(math.log10(abs(v))))
        mant = v / (10 ** exp)
        return (f"{mant:.2f}", exp)
    except (ValueError, TypeError):
        return ("", 0)

def render_header(cell, header_text):
    cell.text = ""
    p = cell.paragraphs[0]
    unit_idx = header_text.find(" (")
    name = header_text if unit_idx == -1 else header_text[:unit_idx]
    unit = "" if unit_idx == -1 else header_text[unit_idx:]
    if "_" in name:
        base, sub = name.split("_", 1)
        p.add_run(base).bold = True
        p.add_run(sub).bold = True.font.subscript = True
        if unit: p.add_run(unit).bold = True
    else:
        p.add_run(header_text).bold = True

def build_formatted_df(df):
    required = ["index", "t_start_min", "t_end_min", "rt_min", "apex_intensity", "area", "pct_area"]
    if not all(c in df.columns for c in required): return pd.DataFrame()
    out = pd.DataFrame()
    out["Peak #"] = df["index"].apply(lambda x: str(int(round(float(x)))) if pd.notna(x) else "")
    out["t_start (min)"] = df["t_start_min"].apply(lambda x: f"{float(x):.2f}" if pd.notna(x) else "")
    out["t_end (min)"] = df["t_end_min"].apply(lambda x: f"{float(x):.2f}" if pd.notna(x) else "")
    out["rt (min)"] = df["rt_min"].apply(lambda x: f"{float(x):.2f}" if pd.notna(x) else "")
    out["Height_max (AU)"] = df["apex_intensity"].apply(lambda x: f"{float(x):.3f}" if pd.notna(x) else "")
    out["Area (AU)"] = df["area"]
    out["Area (%)"] = df["pct_area"].apply(lambda x: f"{float(x):.2f} %" if pd.notna(x) else "")
    return out

def add_peaks_table_to_docx(doc, df_raw):
    if "measure_value" in df_raw.columns:
        df_raw = df_raw.drop(columns=["measure_value"])
    df_fmt = build_formatted_df(df_raw)
    if df_fmt.empty:
        doc.add_paragraph("No peaks found or data format is incorrect.")
        return

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)
    table = doc.add_table(rows=df_fmt.shape[0] + 1, cols=df_fmt.shape[1])
    table.autofit = True
    for j, col in enumerate(df_fmt.columns):
        render_header(table.cell(0, j), str(col))
    for i in range(df_fmt.shape[0]):
        for j, col in enumerate(df_fmt.columns):
            cell = table.cell(i + 1, j)
            p = cell.paragraphs[0]
            if col == "Area (AU)":
                mant, exp = decompose_sci(df_raw["area"].iloc[i])
                if mant:
                    p.add_run(f"{mant}x10")
                    p.add_run(str(exp)).font.superscript = True
            else:
                p.add_run(str(df_fmt.iat[i, j]))

def generate_combined_docx(docx_path: Path, experiment_name: str, replicate_results: list, summary_stats: dict):
    doc = Document()
    set_page_a4(doc.sections[0])
    doc.add_heading(f"Analysis Report for {experiment_name}", level=1)
    doc.add_paragraph(f"Generated on: {time.strftime('%Y-%m-%d %H:%M:%S')}")

    for i, res in enumerate(replicate_results):
        doc.add_heading(f"Replicate {i+1}: {res['name']}", level=2)
        add_top_image(doc, res['plot_disk_path'], doc.sections[0], half_ratio=0.4)
        doc.add_paragraph() # Spacer
        add_peaks_table_to_docx(doc, res['peaks_df'])
        doc.add_page_break()

    if summary_stats:
        doc.add_heading("Summary Statistics for Largest Peak", level=2)
        doc.add_paragraph(f"Statistics for the peak with retention time near {summary_stats['target_rt']:.2f} min.")
        table = doc.add_table(rows=2, cols=4)
        table.style = 'Table Grid'
        headers = ["Target RT (min)", "Avg. Area %", "Std. Dev. (Area %)", "Replicates Found"]
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
        table.cell(1, 0).text = f"{summary_stats['target_rt']:.3f}"
        table.cell(1, 1).text = f"{summary_stats['avg_pct_area']:.3f}"
        table.cell(1, 2).text = f"{summary_stats['std_dev_pct_area']:.3f}"
        table.cell(1, 3).text = str(summary_stats['n_replicates'])

    doc.save(docx_path)

def run_analysis_pipeline(uploaded_file_paths: list[Path], experiment_name: str, options: dict) -> dict:
    """
    Orchestrates the analysis pipeline for multiple replicate files.
    Args:
        uploaded_file_paths (list[Path]): List of paths to user-uploaded .cdf files.
        experiment_name (str): Name for the experiment.
        options (dict): A dictionary of options selected by the user in the UI.
    Returns:
        A dictionary containing the URL to the HTML report and the run ID for downloads.
    """
    original_cwd = Path.cwd()
    run_id = str(uuid.uuid4())[:8]
    run_dir = REPORTS_DIR / f"{run_id}_{experiment_name.replace(' ', '_')}"
    run_dir.mkdir()

    replicate_results = []
    py_executable = sys.executable or shutil.which('python3') or 'python3'

    try:
        print(f"--- Starting pipeline for experiment '{experiment_name}' (Run ID: {run_id}) ---")
        for i, file_path in enumerate(uploaded_file_paths):
            rep_name = file_path.name
            print(f"\n--- Processing replicate {i+1}/{len(uploaded_file_paths)}: {rep_name} ---")
            
            base_name = Path(rep_name.split('_', 1)[-1]).stem
            rep_dir = run_dir / f"replicate_{i+1}"
            rep_dir.mkdir()
            
            target_cdf_path = rep_dir / rep_name
            shutil.copy(file_path, target_cdf_path)
            os.chdir(rep_dir)

            # --- 1. Run cdf14.py ---
            print("Step 1: Running cdf14.py...")
            cmd1 = [py_executable, str(SCRIPTS_DIR_PATH / 'cdf14.py'), '-i', rep_name]
            subprocess.run(cmd1, check=True, capture_output=True, text=True)

            # --- 2. Run cdf13_post3.py ---
            print("Step 2: Running cdf13_post3.py...")
            post_args = {}
            if options.get('preset') in PRESETS:
                post_args.update(PRESETS[options['preset']])
            for key in ['snr', 'smooth', 'min-width-sec', 'noise-mode', 'baseline-method']:
                if options.get(key) is not None:
                    post_args[key] = str(options[key])
            
            post_cmd = [py_executable, str(SCRIPTS_DIR_PATH / 'cdf13_post3.py'), rep_name]
            for key, value in post_args.items():
                post_cmd.extend([f'--{key}', str(value)])
            
            print(f"  > Command: {' '.join(shlex.quote(s) for s in post_cmd)}")
            subprocess.run(post_cmd, check=True, capture_output=True, text=True)

            # --- 3. Collect results ---
            peaks_csv_path = rep_dir / f"{base_name}_peaks.csv"
            plot_png_path = rep_dir / f"{base_name}_with_baseline.png"
            if not (peaks_csv_path.exists() and plot_png_path.exists()):
                raise FileNotFoundError(f"Expected outputs not found for {rep_name}")

            df_peaks = pd.read_csv(peaks_csv_path)
            asset_image_filename = f"plot_{base_name}_{run_id}_{i+1}.png"
            asset_dest_path = ASSETS_DIR / asset_image_filename
            shutil.copy(plot_png_path, asset_dest_path)

            replicate_results.append({
                "name": rep_name,
                "peaks_df": df_peaks,
                "plot_asset_url": f"/assets/{asset_image_filename}",
                "plot_disk_path": asset_dest_path,
            })
            os.chdir(original_cwd)

        # --- 4. Aggregation and Summary Statistics ---
        print("\n--- Aggregating results and calculating statistics ---")
        summary_stats = None
        all_peaks_dfs = [res['peaks_df'] for res in replicate_results if not res['peaks_df'].empty]
        if all_peaks_dfs:
            full_peaks_df = pd.concat(all_peaks_dfs, ignore_index=True)
            if 'area' in full_peaks_df.columns and not full_peaks_df.empty:
                largest_peak = full_peaks_df.loc[full_peaks_df['area'].idxmax()]
                target_rt = largest_peak['rt_min']
                rt_tolerance = 0.1  # minutes

                corresponding_pct_areas = []
                for df in all_peaks_dfs:
                    candidates = df[np.isclose(df['rt_min'], target_rt, atol=rt_tolerance)]
                    if not candidates.empty:
                        best_match = candidates.loc[candidates['area'].idxmax()]
                        corresponding_pct_areas.append(best_match['pct_area'])
                    else:
                        corresponding_pct_areas.append(np.nan)
                
                pct_areas_arr = np.array(corresponding_pct_areas)
                summary_stats = {
                    "target_rt": target_rt,
                    "avg_pct_area": np.nanmean(pct_areas_arr),
                    "std_dev_pct_area": np.nanstd(pct_areas_arr),
                    "n_replicates": len(pct_areas_arr) - np.isnan(pct_areas_arr).sum()
                }
                print(f"Summary for peak at ~{target_rt:.2f} min: Avg Area % = {summary_stats['avg_pct_area']:.2f}, SD = {summary_stats['std_dev_pct_area']:.2f}")

        # --- 5. Generate Combined HTML Report ---
        print("Step 5: Generating combined HTML report...")
        report_html_parts = []
        for i, res in enumerate(replicate_results):
            table_html = res['peaks_df'].to_html(classes='dataframe', border=0, index=False)
            plot_element = f'<img src="{res["plot_asset_url"]}" alt="Plot for {res["name"]}" style="width:100%;">'
            report_html_parts.append({
                "title": f"Replicate {i+1}: {res['name']}",
                "plot": plot_element,
                "table": table_html
            })

        summary_table_html = ""
        if summary_stats:
            summary_df = pd.DataFrame([summary_stats])
            summary_df = summary_df.rename(columns={
                "target_rt": "Target RT (min)",
                "avg_pct_area": "Average Area %",
                "std_dev_pct_area": "Std. Dev. (Area %)",
                "n_replicates": "Replicates Found"
            })
            summary_table_html = summary_df.to_html(classes='dataframe summary-table', border=0, index=False, float_format='%.3f')

        docx_filename = f"{experiment_name.replace(' ', '_')}_report.docx"
        download_url = f"/download/docx/{run_dir.name}/{docx_filename}"
        download_section_html = f'<hr><a href="{download_url}" download><button style="width:100%; padding: 15px; font-size: 16px; cursor: pointer;">Download Combined DOCX Report</button></a>'

        # Render the Jinja2 template
        env = Environment(loader=FileSystemLoader(PROJECT_ROOT))
        template = env.get_template('report_template.html')
        report_html_string = template.render(
            report_title=f"Analysis Report for {experiment_name}",
            replicate_reports=report_html_parts,
            summary_table=summary_table_html,
            download_section=download_section_html,
            generation_date=time.strftime("%Y-%m-%d %H:%M:%S")
        )

        html_report_filename = f"report_{experiment_name.replace(' ', '_')}_{run_id}.html"
        with open(ASSETS_DIR / html_report_filename, "w") as f:
            f.write(report_html_string)

        # --- 6. Generate Combined DOCX Report ---
        print("Step 6: Generating combined DOCX report...")
        docx_report_path = run_dir / docx_filename
        generate_combined_docx(docx_report_path, experiment_name, replicate_results, summary_stats)

        print(f"--- Pipeline finished for {experiment_name} ---")

        return {
            "html_report_url": f"/assets/{html_report_filename}",
            "run_id": run_dir.name # Pass the full run dir name for the download route
        }

    except subprocess.CalledProcessError as e:
        # If a script fails, capture its output to show the user
        error_message = (
            f"A script failed with exit code {e.returncode}.\n\n"
            f"Command:\n{' '.join(shlex.quote(s) for s in e.cmd)}\n\n"
            f"--- STDOUT ---\n{e.stdout}\n\n"
            f"--- STDERR ---\n{e.stderr}"
        )
        print(error_message) # Log the full error to the server console
        # Re-raise the exception to be caught by the Dash callback
        raise Exception(error_message)

    finally:
        # Always change back to the original directory
        os.chdir(original_cwd)
