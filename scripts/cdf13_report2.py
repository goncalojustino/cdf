#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
cdf13_report.py — build a one-page A4 DOCX with plot on top and formatted integration table below.

Usage
  python cdf13_report.py kh18.cdf
  python cdf13_report.py kh18.cdf --img kh18_with_baseline.png --peaks kh18_peaks.csv --out kh18_report.docx

Inputs inferred from <base>=basename(.cdf) unless overridden:
  Figure: <base>_with_baseline.png
  Table : <base>_peaks.csv

DOCX table formatting:
  Headers:
    index            → Peak #
    t_start_min      → t_start (min)      [“start” subscript]
    t_end_min        → t_end (min)        [“end” subscript]
    rt_min           → rt (min)
    apex_intensity   → Height_max (AU)    [“max” subscript]
    area             → Area (AU)
    pct_area         → Area (%)
  Values:
    Peak #                     integer
    t_start, t_end, rt         2 decimals
    Height_max (AU)            3 decimals
    Area (AU)                  scientific, 2 decimals: A.AA x10 with exponent as superscript (e.g., 1.23x10⁻¹)
    Area (%)                   2 decimals + “ %”
    “measure_value” column is omitted from the DOCX.
Deps
  pip install python-docx pandas pillow
"""

import argparse, os, sys, math
import pandas as pd
from PIL import Image
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------- helpers ----------

def infer_paths(cdf_path, img_arg, peaks_arg, out_arg):
    base = os.path.splitext(os.path.basename(cdf_path))[0]
    img = img_arg or f"{base}_with_baseline.png"
    peaks = peaks_arg or f"{base}_peaks.csv"
    out = out_arg or f"{base}_report.docx"
    return img, peaks, out

def ensure_exists(path, kind):
    if not os.path.exists(path):
        sys.exit(f"ERROR: {kind} not found: {path}")

def set_page_a4(section, margin_mm=20):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    m = Mm(margin_mm)
    section.left_margin = m
    section.right_margin = m
    section.top_margin = m
    section.bottom_margin = m

def add_top_image(doc, img_path, section, half_ratio=0.5):
    usable_w = section.page_width - section.left_margin - section.right_margin
    usable_h = section.page_height - section.top_margin - section.bottom_margin
    target_h = usable_h * half_ratio
    with Image.open(img_path) as im:
        w_px, h_px = im.size
    scaled_w_by_h = (w_px / h_px) * target_h
    if scaled_w_by_h <= usable_w:
        doc.add_picture(img_path, height=target_h)
    else:
        doc.add_picture(img_path, width=usable_w)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def fmt_int(x):
    try:
        return str(int(round(float(x))))
    except Exception:
        return ""

def fmt_dec(x, nd=2):
    try:
        return f"{float(x):.{nd}f}"
    except Exception:
        return ""

def decompose_sci(val):
    """Return (mantissa_str with 2 dp, exponent_int)."""
    try:
        v = float(val)
    except Exception:
        return ("", 0)
    if v == 0.0:
        return ("0.00", 0)
    exp = int(math.floor(math.log10(abs(v))))
    mant = v / (10 ** exp)
    return (f"{mant:.2f}", exp)

def fmt_pct2(x):
    try:
        return f"{float(x):.2f} %"
    except Exception:
        return ""

def build_formatted_df(df):
    """
    Input columns expected from cdf13_post.py:
      index, t_start_min, t_end_min, rt_min, apex_intensity, area, pct_area, [measure_value]
    Output columns and order for DOCX:
      Peak #, t_start (min), t_end (min), rt (min), Height_max (AU), Area (AU), Area (%)
    """
    # Validate required columns
    required = ["index","t_start_min","t_end_min","rt_min","apex_intensity","area","pct_area"]
    missing = [c for c in required if c not in df.columns]
    if missing:
        sys.exit(f"ERROR: required columns missing in peaks CSV: {missing}")

    out = pd.DataFrame()
    out["Peak #"] = df["index"].apply(fmt_int)
    out["t_start (min)"] = df["t_start_min"].apply(lambda v: fmt_dec(v, 2))
    out["t_end (min)"] = df["t_end_min"].apply(lambda v: fmt_dec(v, 2))
    out["rt (min)"] = df["rt_min"].apply(lambda v: fmt_dec(v, 2))
    out["Height_max (AU)"] = df["apex_intensity"].apply(lambda v: fmt_dec(v, 3))
    # Area handled specially at render time for superscript. Keep a placeholder.
    out["Area (AU)"] = df["area"]  # keep numeric for rendering
    out["Area (%)"] = df["pct_area"].apply(fmt_pct2)
    return out

def render_header(cell, header_text):
    """Write header with subscript for 't_start', 't_end', 'Height_max', keeping unit suffix."""
    cell.text = ""
    p = cell.paragraphs[0]

    # split unit suffix like " (min)" or " (AU)"
    unit_idx = header_text.find(" (")
    name = header_text if unit_idx == -1 else header_text[:unit_idx]
    unit = "" if unit_idx == -1 else header_text[unit_idx:]

    # subscript part after underscore, if present
    if "_" in name:
        base, sub = name.split("_", 1)
        r1 = p.add_run(base); r1.bold = True
        r2 = p.add_run(sub);  r2.bold = True; r2.font.subscript = True
        if unit:
            r3 = p.add_run(unit); r3.bold = True
    else:
        r = p.add_run(header_text); r.bold = True


def add_table(doc, df_fmt, area_raw_series):
    # Style
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    rows, cols = df_fmt.shape
    table = doc.add_table(rows=rows + 1, cols=cols)
    table.autofit = True

    # Headers with subscript logic
    for j, col in enumerate(df_fmt.columns):
        render_header(table.cell(0, j), str(col))

    # Body with superscript for Area (AU)
    for i in range(rows):
        for j, col in enumerate(df_fmt.columns):
            cell = table.cell(i + 1, j)
            cell.text = ""  # clear default paragraph
            p = cell.paragraphs[0]
            if col == "Area (AU)":
                mant, exp = decompose_sci(area_raw_series.iloc[i])
                r1 = p.add_run(f"{mant}x10")
                r2 = p.add_run(f"{exp}")
                r2.font.superscript = True
            else:
                val = df_fmt.iat[i, j]
                p.add_run("" if pd.isna(val) else str(val))

# ---------- main ----------

def main():
    ap = argparse.ArgumentParser(description="Create A4 DOCX with plot on top and formatted integration table below.")
    ap.add_argument("cdf", help="Original .cdf filename (used only for base name inference)")
    ap.add_argument("--img", help="Figure path; default <base>_with_baseline.png")
    ap.add_argument("--peaks", help="Peaks CSV path; default <base>_peaks.csv")
    ap.add_argument("--out", help="Output DOCX; default <base>_report.docx")
    ap.add_argument("--margin-mm", type=float, default=20.0, help="Page margins in mm; default 20")
    ap.add_argument("--chrom-lw", type=float, help="Chromatogram line width (ignored, for compatibility)")
    args = ap.parse_args()

    img_path, peaks_path, out_path = infer_paths(args.cdf, args.img, args.peaks, args.out)
    ensure_exists(img_path, "figure")
    ensure_exists(peaks_path, "peaks CSV")

    raw_df = pd.read_csv(peaks_path)
    # Drop 'measure_value' if present
    if "measure_value" in raw_df.columns:
        raw_df = raw_df.drop(columns=["measure_value"])

    df_fmt = build_formatted_df(raw_df)

    doc = Document()
    set_page_a4(doc.sections[0], margin_mm=args.margin_mm)

    # Top half: image
    add_top_image(doc, img_path, doc.sections[0], half_ratio=0.5)

    # Spacer
    doc.add_paragraph("")

    # Bottom half: table
    add_table(doc, df_fmt, area_raw_series=raw_df["area"])

    doc.save(out_path)
    print(f"Wrote: {out_path}")

if __name__ == "__main__":
    main()
